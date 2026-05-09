"""
BookWise AI — Document Processing
Handles PDF, DOCX, TXT, and other document formats with text extraction.
"""
import os
import json
from typing import Tuple, Optional
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from markdownify import markdownify as md
from bs4 import BeautifulSoup

# ═══════════════════════════════════════════════════════════════
#  DOCUMENT EXTRACTION
# ═══════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_path: str) -> Tuple[str, int]:
    """
    Extract text from PDF file.
    
    Returns:
        (text, page_count)
    """
    try:
        reader = PdfReader(file_path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
            except Exception as e:
                print(f"Error extracting page {page_num}: {e}")
        
        full_text = "\n\n".join(text_parts)
        return full_text, len(reader.pages)
    except Exception as e:
        raise Exception(f"Error reading PDF: {e}")


def extract_text_from_docx(file_path: str) -> Tuple[str, int]:
    """
    Extract text from DOCX file.
    
    Returns:
        (text, paragraph_count)
    """
    try:
        doc = DocxDocument(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                paragraphs.append("\n".join(table_text))
        
        full_text = "\n".join(paragraphs)
        return full_text, len(paragraphs)
    except Exception as e:
        raise Exception(f"Error reading DOCX: {e}")


def extract_text_from_txt(file_path: str) -> Tuple[str, int]:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        line_count = len(content.split('\n'))
        return content, line_count
    except Exception as e:
        raise Exception(f"Error reading TXT: {e}")


def extract_text_from_html(file_path: str) -> Tuple[str, int]:
    """Extract text from HTML file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove scripts and styles
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text(separator='\n', strip=True)
        line_count = len(text.split('\n'))
        
        return text, line_count
    except Exception as e:
        raise Exception(f"Error reading HTML: {e}")


def extract_text_from_json(file_path: str) -> Tuple[str, int]:
    """Extract text from JSON file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        def flatten_json(obj, prefix=''):
            result = []
            if isinstance(obj, dict):
                for key, value in obj.items():
                    result.extend(flatten_json(value, f"{prefix}{key}: "))
            elif isinstance(obj, list):
                for item in obj:
                    result.extend(flatten_json(item, prefix))
            else:
                result.append(f"{prefix}{obj}".strip())
            return result
        
        lines = flatten_json(data)
        text = "\n".join(filter(None, lines))
        
        return text, len(lines)
    except Exception as e:
        raise Exception(f"Error reading JSON: {e}")


# ═══════════════════════════════════════════════════════════════
#  DOCUMENT PROCESSOR
# ═══════════════════════════════════════════════════════════════

class DocumentProcessor:
    """Universal document processor"""
    
    SUPPORTED_FORMATS = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_docx,
        '.txt': extract_text_from_txt,
        '.html': extract_text_from_html,
        '.htm': extract_text_from_html,
        '.json': extract_text_from_json,
        '.md': extract_text_from_txt,  # Markdown is just text
        '.xml': extract_text_from_txt,
        '.csv': extract_text_from_txt,
    }
    
    @staticmethod
    def get_supported_formats():
        """Get list of supported file extensions"""
        return list(DocumentProcessor.SUPPORTED_FORMATS.keys())
    
    @staticmethod
    def process(file_path: str) -> Tuple[str, dict]:
        """
        Process any supported document.
        
        Returns:
            (extracted_text, metadata)
        
        Raises:
            ValueError if format not supported
        """
        # Get file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext not in DocumentProcessor.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {ext}. Supported: {DocumentProcessor.get_supported_formats()}")
        
        # Get file metadata
        file_size = os.path.getsize(file_path)
        filename = os.path.basename(file_path)
        
        # Extract text
        extractor = DocumentProcessor.SUPPORTED_FORMATS[ext]
        text, page_count = extractor(file_path)
        
        # Build metadata
        metadata = {
            "filename": filename,
            "file_type": ext[1:],  # Remove leading dot
            "file_size": file_size,
            "char_count": len(text),
            "page_count": page_count,
            "word_count": len(text.split()),
        }
        
        return text, metadata
    
    @staticmethod
    def generate_summary(text: str, max_length: int = 500) -> str:
        """
        Generate a simple summary using extractive summarization.
        
        Note: For better summaries, use Gemini API.
        """
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        
        if len(sentences) <= 2:
            return text[:max_length]
        
        # Score sentences by keyword frequency
        from collections import Counter
        words = ' '.join(sentences).lower().split()
        word_freq = Counter(words)
        
        # Filter out common words
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'is', 'are', 'be', 'was', 'were', 'of', 'with', 'by', 'from', 'up', 'about', 'into', 'through', 'during'}
        word_freq = {w: f for w, f in word_freq.items() if w not in stop_words and len(w) > 3}
        
        sentence_scores = {}
        for i, sent in enumerate(sentences):
            for word in sent.lower().split():
                if word in word_freq:
                    sentence_scores[i] = sentence_scores.get(i, 0) + word_freq[word]
        
        # Get top sentences
        top_sentences = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)[:3]
        top_sentences = sorted(top_sentences, key=lambda x: x[0])  # Sort by original position
        
        summary = '. '.join([sentences[i] for i, _ in top_sentences if i < len(sentences)])
        
        return (summary[:max_length] + '...') if len(summary) > max_length else summary


# ═══════════════════════════════════════════════════════════════
#  DOCUMENT ANALYSIS
# ═══════════════════════════════════════════════════════════════

def analyze_document(text: str) -> dict:
    """Analyze document content"""
    lines = text.split('\n')
    words = text.split()
    sentences = [s.strip() for s in text.split('.') if s.strip()]
    
    # Reading time estimation (average 200 words per minute)
    reading_time_minutes = max(1, len(words) // 200)
    
    # Extract key phrases
    from collections import Counter
    word_freq = Counter(words)
    common_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'is', 'are', 'be', 'was', 'were'}
    key_words = [w for w, _ in word_freq.most_common(10) if w.lower() not in common_words and len(w) > 3]
    
    return {
        "char_count": len(text),
        "word_count": len(words),
        "sentence_count": len(sentences),
        "line_count": len(lines),
        "avg_sentence_length": len(words) // max(1, len(sentences)),
        "avg_word_length": sum(len(w) for w in words) / max(1, len(words)),
        "reading_time_minutes": reading_time_minutes,
        "key_words": key_words,
    }


def extract_quiz_questions(text: str, count: int = 5) -> list:
    """Generate simple quiz questions from document"""
    sentences = [s.strip() for s in text.split('.') if len(s.split()) > 5]
    
    if len(sentences) < count:
        return []
    
    import random
    selected = random.sample(sentences, min(count, len(sentences)))
    
    questions = []
    for sent in selected:
        # Create simple fill-in-the-blank questions
        words = sent.split()
        if len(words) > 5:
            blank_idx = random.randint(1, len(words) - 2)
            answer = words[blank_idx]
            question_text = ' '.join(words[:blank_idx] + ['_____'] + words[blank_idx + 1:])
            questions.append({
                "question": question_text,
                "answer": answer,
            })
    
    return questions
